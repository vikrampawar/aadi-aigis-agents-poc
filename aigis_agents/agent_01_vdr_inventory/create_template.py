"""
One-time script to create the branded DRL Word template.
Run: python -m aigis_agents.agent_01_vdr_inventory.create_template
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_page_number(paragraph):
    """Insert page number field into paragraph."""
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar2)


def create_drl_template(output_path: Path):
    """Create the branded DRL Word template with Aigis styling."""
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ── Define styles ─────────────────────────────────────────────────────────
    styles = doc.styles

    # Heading 1 — Section title (dark navy)
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

    # Heading 2 — Category header
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(11)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    # Normal text
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # ── Cover Page ────────────────────────────────────────────────────────────
    # Aigis header bar (dark navy table row)
    header_table = doc.add_table(rows=1, cols=1)
    header_table.style = "Table Grid"
    cell = header_table.cell(0, 0)
    _set_cell_bg(cell, "0F4C81")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("  AIGIS ANALYTICS  —  CONFIDENTIAL")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph("")

    # Deal name placeholder (Jinja2 tag)
    p_project = doc.add_paragraph()
    p_project.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_project.add_run("Project  {{ deal_name | upper }}")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

    p_title = doc.add_paragraph()
    run = p_title.add_run("DATA REQUEST LIST — ROUND {{ round_number }}")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x6F, 0xBA)

    doc.add_paragraph("")

    # Metadata table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_data = [
        ("Prepared by", "{{ buyer_name }}"),
        ("Date", "{{ run_date }}"),
        ("Deal Type", "{{ deal_type_label }}"),
        ("Jurisdiction", "{{ jurisdiction }}"),
    ]
    for i, (label, value) in enumerate(meta_data):
        row = meta_table.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        _set_cell_bg(lc, "EFF6FF")
        lp = lc.paragraphs[0]
        lp.add_run(label).font.bold = True
        lp.runs[0].font.name = "Calibri"
        lp.runs[0].font.size = Pt(10)
        vp = vc.paragraphs[0]
        vp.add_run(value)
        vp.runs[0].font.name = "Calibri"
        vp.runs[0].font.size = Pt(10)

    doc.add_paragraph("")
    p_conf = doc.add_paragraph()
    run = p_conf.add_run("STRICTLY PRIVATE AND CONFIDENTIAL")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    # ── Summary block ─────────────────────────────────────────────────────────
    doc.add_page_break()

    p_sum = doc.add_paragraph()
    run = p_sum.add_run("EXECUTIVE SUMMARY")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "This Data Request List has been prepared by {{ buyer_name }} in connection with "
        "the proposed acquisition of {{ deal_name }} (the \"Transaction\"). "
        "The requests below are based on a review of the Virtual Data Room as at {{ run_date }}.\n\n"
        "Checklist version: {{ checklist_version }} | "
        "Total VDR files reviewed: {{ total_files }} | "
        "Critical gaps (Need to Have): {{ missing_nth }} missing, {{ partial_nth }} partial | "
        "Supplementary gaps (Good to Have): {{ missing_gth }} missing, {{ partial_gth }} partial"
    )
    p_intro.runs[0].font.name = "Calibri"
    p_intro.runs[0].font.size = Pt(10)

    # ── Section A: Critical Data Requests ────────────────────────────────────
    doc.add_paragraph("")
    p_sec_a = doc.add_heading("SECTION A — CRITICAL DATA REQUESTS (Need to Have)", level=1)

    p_sec_a_intro = doc.add_paragraph()
    p_sec_a_intro.add_run(
        "The following items are classified as Need to Have for a {{ deal_type_label }} transaction "
        "in {{ jurisdiction }}. These should be provided as a priority."
    )
    p_sec_a_intro.runs[0].font.name = "Calibri"
    p_sec_a_intro.runs[0].font.size = Pt(10)
    p_sec_a_intro.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    doc.add_paragraph("{% for item in critical_items %}")

    p_cat = doc.add_heading("{{ item.category_label }}", level=2)

    p_item = doc.add_paragraph()
    p_item.style = "List Number"
    run = p_item.add_run("{{ item.description }}")
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)

    p_status = doc.add_paragraph()
    p_status.add_run("Status: {{ item.status_label }}  |  Item ref: {{ item.id }}")
    p_status.runs[0].font.name = "Calibri"
    p_status.runs[0].font.size = Pt(9)
    p_status.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    p_request = doc.add_paragraph()
    p_request.add_run("Request: {{ item.drl_request_text }}")
    p_request.runs[0].font.name = "Calibri"
    p_request.runs[0].font.size = Pt(10)

    p_note = doc.add_paragraph()
    p_note.add_run("{% if item.notes %}Note: {{ item.notes }}{% endif %}")
    p_note.runs[0].font.name = "Calibri"
    p_note.runs[0].font.size = Pt(9)
    p_note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("{% endfor %}")

    # ── Section B: Supplementary Data Requests ────────────────────────────────
    doc.add_paragraph("")
    p_sec_b = doc.add_heading("SECTION B — SUPPLEMENTARY DATA REQUESTS (Good to Have)", level=1)

    p_sec_b_intro = doc.add_paragraph()
    p_sec_b_intro.add_run(
        "The following items are classified as Good to Have. While not critical to completing "
        "due diligence, these documents would strengthen our analysis."
    )
    p_sec_b_intro.runs[0].font.name = "Calibri"
    p_sec_b_intro.runs[0].font.size = Pt(10)

    doc.add_paragraph("{% for item in supplementary_items %}")

    p_cat_b = doc.add_heading("{{ item.category_label }}", level=2)

    p_item_b = doc.add_paragraph()
    p_item_b.style = "List Number"
    run_b = p_item_b.add_run("{{ item.description }}")
    run_b.font.bold = True
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)

    p_request_b = doc.add_paragraph()
    p_request_b.add_run("Request: {{ item.drl_request_text }}")
    p_request_b.runs[0].font.name = "Calibri"
    p_request_b.runs[0].font.size = Pt(10)

    doc.add_paragraph("{% endfor %}")

    # ── Footer ────────────────────────────────────────────────────────────────
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run("Generated by Aigis Analytics  |  aigis.claudecodecentrallondon.com  |  Strictly Confidential  |  Page ")
    run_f.font.name = "Calibri"
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    _add_page_number(fp)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"DRL template created: {output_path}")


if __name__ == "__main__":
    template_path = Path(__file__).parent.parent.parent / "checklists" / "templates" / "drl_template.docx"
    create_drl_template(template_path)
