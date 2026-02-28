"""
DOCX Data Request List (DRL) generator.

Uses docxtpl (Jinja2-based Word templating) with drl_template.docx.
Falls back to programmatic generation if template not found.
"""

from __future__ import annotations

from pathlib import Path

from aigis_agents.agent_01_vdr_inventory.models import (
    ChecklistItemResult,
    ChecklistStatus,
    DocumentTier,
    GapReport,
)

TEMPLATE_PATH = Path(__file__).parent.parent.parent / "checklists" / "templates" / "drl_template.docx"

_DEAL_TYPE_LABELS = {
    "producing_asset": "Producing Asset Acquisition",
    "exploration": "Exploration Asset Acquisition",
    "development": "Development Asset Acquisition",
    "corporate": "Corporate Acquisition",
}

_STATUS_LABELS = {
    ChecklistStatus.missing: "❌ Not provided",
    ChecklistStatus.partial: "⚠️ Incomplete / requires follow-up",
}


def _build_drl_items(items: list[ChecklistItemResult], tier: DocumentTier) -> list[dict]:
    """Build the context dicts for docxtpl from gap report items."""
    drl_items = []
    for item in items:
        if item.tier != tier:
            continue
        if item.status not in (ChecklistStatus.missing, ChecklistStatus.partial):
            continue
        drl_items.append({
            "id": item.item_id,
            "description": item.description,
            "category_label": item.category_label,
            "status_label": _STATUS_LABELS.get(item.status, item.status.value),
            "drl_request_text": item.drl_request_text or f"Please provide {item.description.lower()}.",
            "notes": item.notes,
            "is_partial": item.status == ChecklistStatus.partial,
            "matched_files": [f.filename for f in item.matched_files[:2]],
        })
    return drl_items


def _programmatic_drl(
    gap_report: GapReport,
    buyer_name: str,
    round_number: int,
    output_path: Path,
) -> Path:
    """Fallback: generate DRL using python-docx without a template."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def _set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    run_date = gap_report.run_timestamp[:10]
    deal_type_label = _DEAL_TYPE_LABELS.get(gap_report.deal_type.value, gap_report.deal_type.value)

    # Header bar
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    _set_cell_bg(cell, "0F4C81")
    p = cell.paragraphs[0]
    run = p.add_run("  AIGIS ANALYTICS  —  STRICTLY CONFIDENTIAL")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph("")

    # Title
    p_title = doc.add_paragraph()
    p_title.add_run(f"PROJECT {gap_report.deal_name.upper()}").font.bold = True
    p_title.runs[0].font.name = "Calibri"
    p_title.runs[0].font.size = Pt(20)
    p_title.runs[0].font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

    p_sub = doc.add_paragraph()
    p_sub.add_run(f"Data Request List — Round {round_number}").font.bold = True
    p_sub.runs[0].font.name = "Calibri"
    p_sub.runs[0].font.size = Pt(14)
    p_sub.runs[0].font.color.rgb = RGBColor(0x1A, 0x6F, 0xBA)

    doc.add_paragraph("")

    # Metadata table
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows_data = [
        ("Prepared by", buyer_name or "Aigis Analytics"),
        ("Date", run_date),
        ("Deal Type", deal_type_label),
        ("Jurisdiction", gap_report.jurisdiction.value),
    ]
    for i, (lbl, val) in enumerate(rows_data):
        lc, vc = meta.rows[i].cells
        _set_cell_bg(lc, "EFF6FF")
        lc.paragraphs[0].add_run(lbl).font.bold = True
        lc.paragraphs[0].runs[0].font.name = "Calibri"
        lc.paragraphs[0].runs[0].font.size = Pt(10)
        vc.paragraphs[0].add_run(val)
        vc.paragraphs[0].runs[0].font.name = "Calibri"
        vc.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph("")
    s = gap_report.summary
    p_sum = doc.add_paragraph()
    p_sum.add_run(
        f"Summary: {s.missing_nth} critical (Need to Have) and {s.missing_gth} supplementary "
        f"(Good to Have) documents are not yet in the VDR. "
        f"An additional {s.partial_nth + s.partial_gth} item(s) require follow-up."
    )
    p_sum.runs[0].font.name = "Calibri"
    p_sum.runs[0].font.size = Pt(10)

    doc.add_page_break()

    def _write_section(heading: str, items: list[dict], intro: str, color: str):
        h = doc.add_paragraph()
        run = h.add_run(heading)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(
            int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        )

        p_intro = doc.add_paragraph()
        p_intro.add_run(intro)
        p_intro.runs[0].font.name = "Calibri"
        p_intro.runs[0].font.size = Pt(10)
        doc.add_paragraph("")

        current_cat = None
        counter = 1
        for item in items:
            if item["category_label"] != current_cat:
                current_cat = item["category_label"]
                h2 = doc.add_paragraph()
                h2.add_run(current_cat).font.bold = True
                h2.runs[0].font.name = "Calibri"
                h2.runs[0].font.size = Pt(11)
                h2.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

            p = doc.add_paragraph(style="List Number")
            run = p.add_run(item["description"])
            run.font.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10)

            p_status = doc.add_paragraph()
            p_status.add_run(f"Status: {item['status_label']}  |  Ref: {item['id']}")
            p_status.runs[0].font.name = "Calibri"
            p_status.runs[0].font.size = Pt(9)
            p_status.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

            p_req = doc.add_paragraph()
            p_req.add_run(f"Request: {item['drl_request_text']}")
            p_req.runs[0].font.name = "Calibri"
            p_req.runs[0].font.size = Pt(10)

            if item.get("notes"):
                p_note = doc.add_paragraph()
                p_note.add_run(f"Note: {item['notes']}")
                p_note.runs[0].font.name = "Calibri"
                p_note.runs[0].font.size = Pt(9)
                p_note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

            if item.get("is_partial") and item.get("matched_files"):
                p_found = doc.add_paragraph()
                p_found.add_run(f"Found in VDR: {', '.join(item['matched_files'])} — please confirm if complete.")
                p_found.runs[0].font.name = "Calibri"
                p_found.runs[0].font.size = Pt(9)
                p_found.runs[0].font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)

            doc.add_paragraph("")
            counter += 1

    critical_items = _build_drl_items(gap_report.items, DocumentTier.need_to_have)
    supplementary_items = _build_drl_items(gap_report.items, DocumentTier.good_to_have)

    if critical_items:
        _write_section(
            "SECTION A — CRITICAL DATA REQUESTS (Need to Have)",
            critical_items,
            f"The following {len(critical_items)} items are classified as Need to Have "
            f"for a {deal_type_label} in {gap_report.jurisdiction.value}. "
            "These should be provided as a priority before submission of any bid.",
            "DC2626",
        )

    if supplementary_items:
        if critical_items:
            doc.add_page_break()
        _write_section(
            "SECTION B — SUPPLEMENTARY DATA REQUESTS (Good to Have)",
            supplementary_items,
            "The following items would strengthen our analysis. "
            "Please provide where available.",
            "0F4C81",
        )

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        f"Generated by Aigis Analytics  |  Checklist {gap_report.checklist_version}  |  {run_date}  |  Strictly Confidential"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def generate_drl(
    gap_report: GapReport,
    output_path: Path,
    buyer_name: str | None = None,
    round_number: int = 1,
) -> Path:
    """
    Generate the Data Request List DOCX.
    Uses docxtpl template if available; falls back to programmatic generation.
    """
    critical_items = _build_drl_items(gap_report.items, DocumentTier.need_to_have)
    supplementary_items = _build_drl_items(gap_report.items, DocumentTier.good_to_have)

    if not critical_items and not supplementary_items:
        # Nothing to request — write a clean "no gaps" DOCX
        from docx import Document
        doc = Document()
        doc.add_heading(f"Project {gap_report.deal_name} — Data Request List", level=1)
        doc.add_paragraph(
            f"No outstanding data requests as at {gap_report.run_timestamp[:10]}. "
            "All checklist items are present in the VDR."
        )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    if TEMPLATE_PATH.exists():
        try:
            from docxtpl import DocxTemplate
            tpl = DocxTemplate(str(TEMPLATE_PATH))
            run_date = gap_report.run_timestamp[:10]
            deal_type_label = _DEAL_TYPE_LABELS.get(gap_report.deal_type.value, gap_report.deal_type.value)
            s = gap_report.summary
            context = {
                "deal_name": gap_report.deal_name,
                "buyer_name": buyer_name or "Aigis Analytics",
                "run_date": run_date,
                "round_number": round_number,
                "deal_type_label": deal_type_label,
                "jurisdiction": gap_report.jurisdiction.value,
                "checklist_version": gap_report.checklist_version,
                "total_files": s.total_files,
                "missing_nth": s.missing_nth,
                "partial_nth": s.partial_nth,
                "missing_gth": s.missing_gth,
                "partial_gth": s.partial_gth,
                "critical_items": critical_items,
                "supplementary_items": supplementary_items,
            }
            tpl.render(context)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            tpl.save(str(output_path))
            return output_path
        except Exception:
            pass  # Fall through to programmatic

    return _programmatic_drl(gap_report, buyer_name or "Aigis Analytics", round_number, output_path)
